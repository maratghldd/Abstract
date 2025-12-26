import os
import re
from pathlib import Path


# Установите зависимости: pip install pdfplumber python-docx

def extract_text_from_file(file_path):
    """
    Извлекает текст из файлов PDF, DOCX, DOC и TXT

    Args:
        file_path: Путь к файлу

    Returns:
        Текст из файла или None в случае ошибки
    """
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        elif file_extension == '.doc':
            return extract_text_from_doc(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        else:
            print(f"❌ Неподдерживаемый формат файла: {file_extension}")
            return None

    except Exception as e:
        print(f"❌ Ошибка при чтении файла {file_path}: {e}")
        return None


# ============= PDF =============

import pdfplumber


def extract_text_from_pdf(pdf_path):
    """Извлекает текст из PDF файла"""
    print(f"📄 Читаю PDF: {pdf_path}")

    try:
        text_parts = []

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Извлекаем текст со страницы
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    # Если текст не извлекается, попробуем OCR или просто пропустим
                    print(f"   Страница {page_num}: текст не найден")

        full_text = '\n\n'.join(text_parts)
        print(f"✅ Прочитано {len(pdf.pages)} страниц, {len(full_text)} символов")
        return full_text

    except Exception as e:
        print(f"❌ Ошибка при чтении PDF: {e}")
        return None


# ============= DOCX =============

from docx import Document


def extract_text_from_docx(docx_path):
    """Извлекает текст из DOCX файла"""
    print(f"📄 Читаю DOCX: {docx_path}")

    try:
        doc = Document(docx_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Также проверяем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = '\n\n'.join(text_parts)
        print(f"✅ Прочитано {len(full_text)} символов")
        return full_text

    except Exception as e:
        print(f"❌ Ошибка при чтении DOCX: {e}")
        return None


# ============= DOC (старый формат) =============

def extract_text_from_doc(doc_path):
    """Извлекает текст из DOC файла (старый формат)"""
    print(f"📄 Читаю DOC: {doc_path}")

    try:
        # Для DOC файлов нужно использовать дополнительные библиотеки
        # Попробуем сначала антиватный textract или python-doc
        import subprocess
        import tempfile

        # Способ 1: Используем antiword (нужно установить)
        try:
            result = subprocess.run(['antiword', doc_path],
                                    capture_output=True, text=True, encoding='utf-8')
            if result.returncode == 0:
                text = result.stdout
                print(f"✅ Прочитано {len(text)} символов (через antiword)")
                return text
        except FileNotFoundError:
            print("⚠️  antiword не установлен, пробую другие методы...")

        # Способ 2: Конвертируем в txt через LibreOffice
        try:
            with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
                tmp_path = tmp.name

            # Конвертируем через LibreOffice
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt',
                            doc_path, '--outdir', os.path.dirname(tmp_path)],
                           capture_output=True)

            txt_file = doc_path.replace('.doc', '.txt')
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                os.remove(txt_file)  # Удаляем временный файл
                print(f"✅ Прочитано {len(text)} символов (через конвертацию)")
                return text
        except Exception as e:
            print(f"⚠️  Ошибка конвертации: {e}")

        print("❌ Не удалось прочитать DOC файл. Установите antiword или LibreOffice")
        return None

    except Exception as e:
        print(f"❌ Ошибка при чтении DOC: {e}")
        return None


# ============= TXT =============

def extract_text_from_txt(txt_path):
    """Читает текст из TXT файла"""
    print(f"📄 Читаю TXT: {txt_path}")

    try:
        # Пробуем разные кодировки
        encodings = ['utf-8', 'cp1251', 'koi8-r', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"✅ Прочитано {len(text)} символов (кодировка: {encoding})")
                return text
            except UnicodeDecodeError:
                continue

        # Если не сработало, читаем с игнорированием ошибок
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        print(f"⚠️  Прочитано с игнорированием ошибок: {len(text)} символов")
        return text

    except Exception as e:
        print(f"❌ Ошибка при чтении TXT: {e}")
        return None


# ============= ИНТЕГРАЦИЯ С ВАШИМ КОДОМ =============

from transformers import T5ForConditionalGeneration, T5Tokenizer
import torch


def generate_title_from_file(file_path):
    """
    Полный пайплайн: файл → текст → название
    """
    # 1. Извлекаем текст из файла
    text = extract_text_from_file(file_path)

    if not text:
        print("❌ Не удалось извлечь текст из файла")
        return None

    # 2. Обрабатываем текст вашей функцией
    model_path = "./models/rut5_base_sum_gazeta"

    try:
        tokenizer = T5Tokenizer.from_pretrained(model_path, local_files_only=True)
        model = T5ForConditionalGeneration.from_pretrained(model_path, local_files_only=True)
        model.to('cpu')

        # Берем достаточно текста для понимания контекста
        context = text[:600]

        # Оптимальный промпт
        prompt = f"Заголовок документа: {context}"

        # Генерация
        inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=400)

        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_length=50,
                min_length=15,
                num_beams=3,
                early_stopping=True,
                repetition_penalty=1.3,
                length_penalty=1.0,
                no_repeat_ngram_size=2
            )

        title = tokenizer.decode(outputs[0], skip_special_tokens=True)
        title = title.replace("Заголовок документа:", "").replace("Заголовок:", "").strip()

        # Берем только первое предложение
        if '.' in title:
            title = title.split('.')[0].strip()

        # До 25 слов максимум
        words = title.split()
        if len(words) > 25:
            title = " ".join(words[:25])

        return title.strip()

    except Exception as e:
        print(f"❌ Ошибка при генерации названия: {e}")
        return None


# ============= ПАКЕТНАЯ ОБРАБОТКА =============

def process_folder(folder_path, output_file=None):
    """
    Обрабатывает все файлы в папке
    """
    import json
    from datetime import datetime

    print(f"📁 Обрабатываю папку: {folder_path}")

    results = []
    supported_extensions = ['.pdf', '.docx', '.doc', '.txt']

    # Собираем все файлы
    all_files = []
    for ext in supported_extensions:
        all_files.extend(Path(folder_path).glob(f"*{ext}"))
        all_files.extend(Path(folder_path).glob(f"*{ext.upper()}"))

    print(f"📊 Найдено файлов: {len(all_files)}")

    for file_path in all_files:
        print(f"\n{'=' * 60}")
        print(f"Обрабатываю: {file_path.name}")

        # Генерируем название
        title = generate_title_from_file(str(file_path))

        if title:
            result = {
                'filename': file_path.name,
                'title': title,
                'path': str(file_path),
                'processed_at': datetime.now().isoformat(),
                'word_count': len(title.split())
            }
            results.append(result)

            print(f"✅ Название: {title}")
            print(f"📏 Слов: {len(title.split())}")
        else:
            print(f"❌ Не удалось обработать файл")

    # Сохраняем результаты
    if output_file and results:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        print(f"\n💾 Результаты сохранены в: {output_file}")

    return results


# ============= ПРИМЕР ИСПОЛЬЗОВАНИЯ =============

if __name__ == "__main__":
    import sys

    print("🎯 СИСТЕМА ОБРАБОТКИ ДОКЛАДОВ")
    print("=" * 60)

    if len(sys.argv) > 1:
        # Если передан путь к файлу или папке
        input_path = sys.argv[1]

        if os.path.isfile(input_path):
            # Обработка одного файла
            title = generate_title_from_file(input_path)
            if title:
                print(f"\n🏷️  Название: {title}")
            else:
                print("❌ Не удалось обработать файл")

        elif os.path.isdir(input_path):
            # Обработка папки
            output_file = "results.json"
            if len(sys.argv) > 2:
                output_file = sys.argv[2]

            process_folder(input_path, output_file)
        else:
            print(f"❌ Путь не найден: {input_path}")
    else:
        # Интерактивный режим
        print("\nВыберите режим:")
        print("1. Обработать один файл")
        print("2. Обработать папку с файлами")
        print("3. Выход")

        choice = input("\nВведите номер (1-3): ").strip()

        if choice == "1":
            file_path = input("Введите путь к файлу: ").strip()
            if os.path.exists(file_path):
                title = generate_title_from_file(file_path)
                if title:
                    print(f"\n🏷️  Название: {title}")
                else:
                    print("❌ Не удалось обработать файл")
            else:
                print("❌ Файл не найден")

        elif choice == "2":
            folder_path = input("Введите путь к папке: ").strip()
            if os.path.isdir(folder_path):
                output_file = input("Имя файла для результатов (по умолчанию results.json): ").strip()
                if not output_file:
                    output_file = "results.json"
                process_folder(folder_path, output_file)
            else:
                print("❌ Папка не найдена")